"""
Phase 2 output rendering: deterministic .docx generation for the weekly
Project Health Report, replacing the previous Markdown output.

Deliberately separate from graph.py / metrics.py, same split as
deck_builder.py: those modules decide *what* the content is (computed
signals + LLM/fallback narrative), this module decides *how it looks* on
the page. No model is ever asked to produce Word XML or fight with layout.

Palette matches deck_builder.py's "Midnight Executive" (navy / ice blue /
white) PMO register, so the weekly Word report and the monthly PowerPoint
deck read as one consistent corporate reporting family.
"""

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# --- Palette (kept identical to deck_builder.py) ----------------------------
NAVY = RGBColor(0x1E, 0x27, 0x61)
NAVY_DEEP = RGBColor(0x12, 0x18, 0x40)
ICE = RGBColor(0xCA, 0xDC, 0xFC)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x6B, 0x6F, 0x7A)
CARD_TINT = "F3F5FB"
LINE_TINT = "E1E5F0"

RAG_COLORS = {
    "Green": RGBColor(0x1E, 0x7D, 0x32),
    "Amber": RGBColor(0xC2, 0x7C, 0x0E),
    "Red": RGBColor(0xB0, 0x28, 0x28),
}
RAG_HEX = {"Green": "1E7D32", "Amber": "C27C0E", "Red": "B02828"}
SCORE_COLORS = {0: RAG_COLORS["Green"], 1: RAG_COLORS["Amber"], 2: RAG_COLORS["Red"]}
SCORE_HEX = {0: "1E7D32", 1: "C27C0E", 2: "B02828"}

ORG_NAME = "Delivery PMO"  # cosmetic letterhead label; change per engagement


# --- low-level OOXML helpers (python-docx has no high-level API for these) --

def _shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_borders(cell, color="E1E5F0", sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def _no_wrap_table_autofit(table):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _bottom_border(paragraph, color="1E2761", sz=18):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_font(run, size=10.5, color=DARK_TEXT, bold=False, italic=False, name="Calibri"):
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name


def _heading(doc, text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    _set_font(run, size=size, color=NAVY, bold=True)
    _bottom_border(p, color="CADCFC", sz=8)
    return p


def _bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_lead:
        r = p.add_run(bold_lead)
        _set_font(r, bold=True)
        text = text[len(bold_lead):]
    r = p.add_run(text)
    _set_font(r)
    return p


def _callout(doc, text, hex_bg, hex_border):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _shade_cell(cell, hex_bg)
    _cell_borders(cell, color=hex_border, sz=6)
    _set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    cell.paragraphs[0].text = ""
    r = cell.paragraphs[0].add_run(text)
    _set_font(r, size=10, color=DARK_TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def _setup_header_footer(section, project_name: str):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    run = hp.add_run(f"Project Health Report  |  {project_name}")
    _set_font(run, size=8.5, color=MUTED)
    hp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_run = fp.add_run(f"{ORG_NAME} — Confidential  |  Page ")
    _set_font(left_run, size=8, color=MUTED)
    _add_page_number_field(fp)
    tail_run = fp.add_run("")
    _set_font(tail_run, size=8, color=MUTED)


def _cover_banner(doc, report: dict):
    band = report["final_rag_status"]
    band_color_hex = RAG_HEX.get(band, "6B6F7A")

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_wrap_table_autofit(table)
    table.columns[0].width = Inches(4.7)
    table.columns[1].width = Inches(2.1)
    left, right = table.rows[0].cells
    table.rows[0].height = Cm(2.6)

    for cell, width in ((left, Inches(4.7)), (right, Inches(2.1))):
        cell.width = width
        _shade_cell(cell, "1E2761")
        _set_cell_margins(cell, top=200, bottom=200, left=220, right=220)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    left.paragraphs[0].text = ""
    p1 = left.paragraphs[0]
    r = p1.add_run("WEEKLY PROJECT HEALTH REPORT")
    _set_font(r, size=11, color=ICE, bold=True)
    p2 = left.add_paragraph()
    r2 = p2.add_run(report["project_name"])
    _set_font(r2, size=16, color=WHITE, bold=True)
    p3 = left.add_paragraph()
    week_ending = report["generated_at_utc"][:10]
    r3 = p3.add_run(f"Week Ending {week_ending}")
    _set_font(r3, size=9.5, color=ICE)

    right.paragraphs[0].text = ""
    rp1 = right.paragraphs[0]
    rp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr1 = rp1.add_run("RAG STATUS")
    _set_font(rr1, size=8.5, color=ICE, bold=True)
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr2 = rp2.add_run(band.upper())
    _set_font(rr2, size=20, color=WHITE, bold=True)
    rp3 = right.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr3 = rp3.add_run(f"Composite Score: {report['composite_score']:.1f} / 100")
    _set_font(rr3, size=9, color=ICE)

    # thin RAG-colored accent bar below the banner
    bar = doc.add_table(rows=1, cols=1)
    _no_wrap_table_autofit(bar)
    bcell = bar.rows[0].cells[0]
    bar.rows[0].height = Cm(0.18)
    _shade_cell(bcell, band_color_hex)
    bcell.paragraphs[0].text = ""


def _signal_table(doc, signal_breakdown: dict):
    headers = ["Signal", "Score", "Detail"]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_wrap_table_autofit(table)
    widths = [Inches(1.9), Inches(0.9), Inches(4.0)]
    for i, w in enumerate(widths):
        table.columns[i].width = w

    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].width = widths[i]
        _shade_cell(hdr_cells[i], "1E2761")
        _cell_borders(hdr_cells[i])
        _set_cell_margins(hdr_cells[i])
        hdr_cells[i].paragraphs[0].text = ""
        r = hdr_cells[i].paragraphs[0].add_run(htext)
        _set_font(r, size=9.5, color=WHITE, bold=True)

    for idx, (name, info) in enumerate(signal_breakdown.items()):
        row = table.add_row().cells
        row[0].width, row[1].width, row[2].width = widths
        bg = CARD_TINT if idx % 2 == 0 else "FFFFFF"
        score = info.get("score")
        score_hex = SCORE_HEX.get(score, "6B6F7A")

        for c in row:
            _shade_cell(c, bg)
            _cell_borders(c)
            _set_cell_margins(c)

        row[0].paragraphs[0].text = ""
        r0 = row[0].paragraphs[0].add_run(name.replace("_", " ").title())
        _set_font(r0, size=9.5, bold=True)

        row[1].paragraphs[0].text = ""
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = row[1].paragraphs[0].add_run(str(score))
        _set_font(r1, size=9.5, color=RGBColor.from_string(score_hex), bold=True)

        row[2].paragraphs[0].text = ""
        r2 = row[2].paragraphs[0].add_run(info.get("detail") or "")
        _set_font(r2, size=9)

    legend = doc.add_paragraph()
    legend.paragraph_format.space_before = Pt(4)
    lr = legend.add_run("Score: 0 = healthy   ·   1 = at risk   ·   2 = critical")
    _set_font(lr, size=8, color=MUTED, italic=True)


def _task_detail_table(doc, rows: list):
    """rows: [{"issue": str, "task_name": str, "owner": str|None,
    "phase_milestone": str|None, "detail": str}, ...]. Renders the specific
    tasks behind this week's status — not just aggregate counts — so a PM
    reading the report knows exactly which task(s) need attention."""
    headers = ["Issue", "Task", "Owner", "Phase / Milestone", "Detail"]
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_wrap_table_autofit(table)
    widths = [Inches(1.15), Inches(1.75), Inches(1.05), Inches(1.25), Inches(1.6)]
    for i, w in enumerate(widths):
        table.columns[i].width = w

    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].width = widths[i]
        _shade_cell(hdr_cells[i], "1E2761")
        _cell_borders(hdr_cells[i])
        _set_cell_margins(hdr_cells[i])
        hdr_cells[i].paragraphs[0].text = ""
        r = hdr_cells[i].paragraphs[0].add_run(htext)
        _set_font(r, size=9.5, color=WHITE, bold=True)

    ISSUE_HEX = {
        "Schedule Slippage": RAG_HEX["Red"],
        "Critical Path": RAG_HEX["Red"],
        "Blocker": RAG_HEX["Amber"],
    }
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, w in enumerate(widths):
            cells[i].width = w
        bg = CARD_TINT if idx % 2 == 0 else "FFFFFF"
        for c in cells:
            _shade_cell(c, bg)
            _cell_borders(c)
            _set_cell_margins(c)

        cells[0].paragraphs[0].text = ""
        r0 = cells[0].paragraphs[0].add_run(row["issue"])
        _set_font(r0, size=9, bold=True, color=RGBColor.from_string(ISSUE_HEX.get(row["issue"], "6B6F7A")))

        cells[1].paragraphs[0].text = ""
        r1 = cells[1].paragraphs[0].add_run(row["task_name"])
        _set_font(r1, size=9.5, bold=True)

        cells[2].paragraphs[0].text = ""
        r2 = cells[2].paragraphs[0].add_run(row.get("owner") or "\u2014")
        _set_font(r2, size=9)

        cells[3].paragraphs[0].text = ""
        r3 = cells[3].paragraphs[0].add_run(row.get("phase_milestone") or "\u2014")
        _set_font(r3, size=9)

        cells[4].paragraphs[0].text = ""
        r4 = cells[4].paragraphs[0].add_run(row["detail"])
        _set_font(r4, size=9)


def _build_task_detail_rows(signal_breakdown: dict) -> list:
    """Flatten the task-level evidence carried on schedule_slippage,
    critical_path_exposure, and blockers into one ordered list of rows for
    _task_detail_table — most severe / most concrete evidence first."""
    rows = []
    sched = signal_breakdown.get("schedule_slippage", {})
    for t in sched.get("slipping_tasks") or []:
        rows.append({
            "issue": "Schedule Slippage",
            "task_name": t.get("task_name", "Unnamed task"),
            "owner": t.get("owner"),
            "phase_milestone": t.get("phase_milestone"),
            "detail": f"{t.get('variance_days'):.0f} day(s) behind baseline"
                      f"{' (' + str(t['status']) + ')' if t.get('status') else ''}.",
        })

    cp = signal_breakdown.get("critical_path_exposure", {})
    seen = {r["task_name"] for r in rows}
    for t in cp.get("at_risk_critical_task_details") or []:
        if t.get("task_name") in seen:
            continue  # already listed under Schedule Slippage — avoid a redundant row
        seen.add(t.get("task_name"))
        rows.append({
            "issue": "Critical Path",
            "task_name": t.get("task_name", "Unnamed task"),
            "owner": t.get("owner"),
            "phase_milestone": t.get("phase_milestone"),
            "detail": f"Zero-float task, {t.get('variance_days'):.0f} day(s) behind — no schedule slack remaining.",
        })

    blockers = signal_breakdown.get("blockers", {})
    for b in blockers.get("blocker_details") or []:
        text = b.get("text") or ""
        rows.append({
            "issue": "Blocker",
            "task_name": b.get("task_name") or "General / PM log",
            "owner": b.get("owner"),
            "phase_milestone": None,
            "detail": text if len(text) <= 90 else text[:87] + "...",
        })

    return rows[:8]  # keep the report to one screen's worth of detail


def build_weekly_docx(report: dict, out_path: str) -> str:
    """Render one weekly report dict (graph.py's finalize output) into a
    corporate-styled .docx at out_path. Returns out_path for convenience."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.7)

    _setup_header_footer(section, report["project_name"])

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK_TEXT

    _cover_banner(doc, report)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # --- Review flags / overrides -------------------------------------------
    disagreement = report.get("source_vs_computed_disagreement")
    if disagreement:
        _callout(
            doc,
            "FLAGGED FOR REVIEW — the source file's own Schedule Health/RAG label is "
            f"\"{disagreement['source_value']}\", but this agent independently computed "
            f"\"{disagreement['computed_value']}\". {disagreement['note']}",
            hex_bg="FBEAEA", hex_border=RAG_HEX["Red"],
        )

    if report.get("model_vs_deterministic_disagreement"):
        _callout(doc, "MODEL CROSS-CHECK — " + report["model_vs_deterministic_disagreement"],
                  hex_bg="FDF3E3", hex_border=RAG_HEX["Amber"])

    if report.get("overrides_triggered"):
        doc.add_paragraph()
        _heading(doc, "Override Rules Triggered", size=11)
        for o in report["overrides_triggered"]:
            _bullet(doc, o)

    # --- Why this status -----------------------------------------------------
    _heading(doc, "Why This Status")
    p = doc.add_paragraph()
    r = p.add_run(report.get("plain_english_reasoning") or "No reasoning available.")
    _set_font(r)

    # --- Top drivers -----------------------------------------------------------
    _heading(doc, "Top Drivers")
    drivers = report.get("top_drivers") or []
    if drivers:
        for d in drivers:
            if ":" in d:
                lead, _, rest = d.partition(":")
                _bullet(doc, d, bold_lead=lead + ":")
            else:
                _bullet(doc, d)
    else:
        _bullet(doc, "No material drivers this week.")

    # --- Signal breakdown table ------------------------------------------------
    _heading(doc, "Signal Breakdown")
    _signal_table(doc, report["signal_breakdown"])

    # --- Task-level risk detail -------------------------------------------------
    # Named tasks behind this week's status — schedule slippage, critical-path
    # exposure, and blockers all now carry task_name/owner/phase evidence
    # (see scoring/metrics.py), so the report can point at exactly which
    # task(s) need attention instead of only an aggregate count.
    task_rows = _build_task_detail_rows(report["signal_breakdown"])
    if task_rows:
        _heading(doc, "Task-Level Risk Detail")
        _task_detail_table(doc, task_rows)

    # --- Recommended actions -----------------------------------------------------
    _heading(doc, "Recommended Actions")
    for a in report.get("recommended_actions") or []:
        _bullet(doc, a)

    # --- Stakeholder sentiment ---------------------------------------------------
    _heading(doc, "Stakeholder Sentiment")
    s = report.get("stakeholder_sentiment", {})
    p = doc.add_paragraph()
    r = p.add_run(f"Score: {s.get('score')} (0 = neutral/positive, 2 = escalation) — ")
    _set_font(r, bold=True)
    r2 = p.add_run(s.get("reasoning") or "")
    _set_font(r2)

    # --- Footer metadata note ------------------------------------------------
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(10)
    _bottom_border(meta, color="E1E5F0", sz=6)
    note = (
        f"Data completeness: {report['data_completeness']*100:.0f}%.  "
        f"Generated by: {report['generated_by']}"
        + (f" (reason: {report['fallback_reason']})" if report.get("fallback_reason") else "")
    )
    r = meta.add_run(note)
    _set_font(r, size=8, color=MUTED, italic=True)

    doc.save(out_path)
    return out_path
